"""Script de uso único para generar la guía técnica del proyecto BioHub."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_header_row(table, headers: list, col_widths: list = None, bg: str = "1F4E79"):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(cell, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)


def code_paragraph(doc, text: str):
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    p.style = "Normal"
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def note_paragraph(doc, text: str, color: str = "FFF3CD"):
    """Add a highlighted note paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.3)
    return p


def section_line(doc):
    doc.add_paragraph()


# ── document ───────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("BioHub — Guía Técnica del Microservicio", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph("Estructura del proyecto, ejecución del servidor y guía de pruebas")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(12)

date_p = doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(9)
date_p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════════
# 1. STACK TECNOLÓGICO
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Stack tecnológico", 1)

stack = [
    ("FastAPI",        "Framework web async para la API REST"),
    ("MongoDB + Motor","Base de datos documental con driver async (Motor)"),
    ("aiokafka",       "Consumer asíncrono del broker Kafka del Grupo 1"),
    ("Redis",          "Caché de consultas frecuentes (con fallback a memoria)"),
    ("Pydantic v2",    "Validación de modelos y schemas"),
    ("pytest + pytest-asyncio", "Suite de tests unitarios sin conexión real a BD"),
    ("python-dotenv",  "Gestión de variables de entorno desde archivo .env"),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
add_header_row(tbl, ["Tecnología", "Uso"], col_widths=[5, 12])
for tech, uso in stack:
    row = tbl.add_row()
    row.cells[0].text = tech
    row.cells[1].text = uso
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 2. ESTRUCTURA DEL PROYECTO
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Estructura del proyecto", 1)

p = doc.add_paragraph("A continuación se muestra el árbol de directorios del proyecto "
                       "(se omite la carpeta ")
p.runs[0].font.size = Pt(10)
r = p.add_run("venv/")
r.font.name = "Courier New"
r.font.size = Pt(10)
p.add_run(" y los archivos de caché de Python).").font.size = Pt(10)

tree = """\
biohub-datamanagement/
│
├── main.py                        ← Punto de entrada FastAPI (lifespan, routers, health)
│
├── config.py                      ← Variables de entorno (Settings con pydantic-settings)
│
├── .env                           ← Variables locales (NO se sube al repo)
├── .env.example                   ← Plantilla de variables de entorno
│
├── requirements.txt               ← Dependencias del proyecto + dependencias de testing
├── pytest.ini                     ← Configuración de pytest (asyncio_mode = auto)
│
├── cache/
│   └── cache.py                   ← Caché Redis con fallback a dict en memoria
│
├── database/
│   ├── connection.py              ← Cliente Motor (MongoDB async), colecciones
│   └── models.py                  ← Modelos Pydantic: AuditEntry, enums, respuestas
│
├── kafka_service/
│   ├── consumer.py                ← Consumer con backoff exponencial y reconexión (ECA-02)
│   └── mock_producer.py           ← Endpoint POST /dev/simulate (solo en development)
│
├── routers/
│   ├── auditoria.py               ← GET /auditoria/* — historial, metadatos, sensibilidad
│   └── aprobacion.py              ← GET/POST /aprobacion/* — flujo de aprobación científica
│
├── services/
│   ├── audit_service.py           ← Lógica de negocio principal (AuditoriaService)
│   ├── approval_service.py        ← Lógica de aprobación
│   └── sensitivity_service.py     ← Clasificación PUBLIC / RESTRICTED / CONFIDENTIAL
│
├── tests/
│   ├── __init__.py
│   └── test_audit_service.py      ← 31 tests unitarios (82 % cobertura en audit_service)
│
├── static/
│   └── index.html                 ← UI mínima montada en /ui
│
└── seed_data.py                   ← Script auxiliar para poblar datos de prueba"""

code_paragraph(doc, tree)
section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 3. FLUJO DE DATOS
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Flujo de datos", 1)

flow_text = doc.add_paragraph()
flow_text.runs  # vacío, usaremos add_run
flow_text.add_run(
    "El microservicio tiene dos canales de entrada y uno de salida:\n\n"
).font.size = Pt(10)

flows = [
    ("Entrada 1 — Kafka (async)",
     "El consumer (kafka_service/consumer.py) escucha el topic registros-biologicos "
     "publicado por el Grupo 1. Por cada mensaje llama a create_audit_entry en audit_service.py, "
     "que inserta una entrada inmutable en la colección audit_entries de MongoDB e invalida el caché."),
    ("Entrada 2 — HTTP POST (aprobación)",
     "El Grupo 3 o el director científico puede cambiar el estado de aprobación de un registro "
     "mediante POST /aprobacion/actualizar. Esto también llama a create_audit_entry con "
     "tipo_cambio=CAMBIO_ESTADO, creando una nueva versión inmutable."),
    ("Salida — HTTP GET (consulta)",
     "El Grupo 3 consulta el historial, metadatos de auditoría y estado de aprobación "
     "a través de los endpoints GET. Todas las respuestas pasan por la capa de caché Redis "
     "antes de ir a MongoDB."),
]

for title_text, desc in flows:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{title_text}: ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 4. CONFIGURACIÓN DEL ENTORNO
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Configuración del entorno", 1)

doc.add_heading("4.1 Variables de entorno (.env)", 2)

p = doc.add_paragraph(
    "Copia el archivo .env.example a .env y completa los valores. "
    "El único campo obligatorio para arrancar en modo desarrollo es MONGODB_URI."
)
p.runs[0].font.size = Pt(10)

code_paragraph(doc, "# MongoDB (obligatorio)")
code_paragraph(doc, "MONGODB_URI=mongodb+srv://<usuario>:<password>@cluster.mongodb.net/")
code_paragraph(doc, "MONGODB_DB_NAME=biohub_db")
code_paragraph(doc, "")
code_paragraph(doc, "# Kafka — en desarrollo se usa el mock, no se necesita broker real")
code_paragraph(doc, "KAFKA_BOOTSTRAP_SERVERS=localhost:9092")
code_paragraph(doc, "KAFKA_TOPIC=registros-biologicos")
code_paragraph(doc, "USE_MOCK_KAFKA=true")
code_paragraph(doc, "")
code_paragraph(doc, "# Redis — opcional, cae en memoria si no está disponible")
code_paragraph(doc, "REDIS_URL=redis://localhost:6379")
code_paragraph(doc, "")
code_paragraph(doc, "# Entorno")
code_paragraph(doc, "ENV=development")

note_paragraph(doc,
    "⚠️  Cuando USE_MOCK_KAFKA=true el consumer de Kafka no arranca. "
    "Para simular eventos usa el endpoint POST /dev/simulate (ver sección 6).",
    "FFF3CD")

section_line(doc)

doc.add_heading("4.2 Instalación de dependencias", 2)

p = doc.add_paragraph(
    "Desde la raíz del proyecto, con el entorno virtual activado:"
)
p.runs[0].font.size = Pt(10)

code_paragraph(doc, "# Crear y activar el entorno virtual (solo la primera vez)")
code_paragraph(doc, "python -m venv venv")
code_paragraph(doc, "")
code_paragraph(doc, "# Windows")
code_paragraph(doc, ".\\venv\\Scripts\\activate")
code_paragraph(doc, "")
code_paragraph(doc, "# macOS / Linux")
code_paragraph(doc, "source venv/bin/activate")
code_paragraph(doc, "")
code_paragraph(doc, "# Instalar todas las dependencias (incluye las de testing)")
code_paragraph(doc, "pip install -r requirements.txt")

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 5. ARRANCAR EL SERVIDOR
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Arrancar el servidor", 1)

code_paragraph(doc, "uvicorn main:app --reload --host 0.0.0.0 --port 8000")

p = doc.add_paragraph(
    "Al iniciar, el servidor ejecuta el lifespan definido en main.py en el siguiente orden:"
)
p.runs[0].font.size = Pt(10)

steps = [
    "Conecta a MongoDB (motor async).",
    "Inicializa la caché (Redis si está disponible, memoria si no).",
    "Si USE_MOCK_KAFKA=false, inicia el consumer Kafka con backoff exponencial.",
    "Registra los routers: /auditoria, /aprobacion y /dev (solo en development).",
    "Monta la UI estática en /ui.",
]
for step in steps:
    p = doc.add_paragraph(style="List Number")
    p.add_run(step).font.size = Pt(10)

p = doc.add_paragraph()
p.add_run("Endpoints de verificación rápida:").bold = True
p.runs[0].font.size = Pt(10)

code_paragraph(doc, "GET  http://localhost:8000/         → nombre, versión, enlace a /docs")
code_paragraph(doc, "GET  http://localhost:8000/health   → estado de MongoDB, caché, Kafka")
code_paragraph(doc, "GET  http://localhost:8000/docs     → Swagger UI interactivo")

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 6. ENDPOINTS DISPONIBLES
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Endpoints disponibles", 1)

endpoints = [
    ("GET",  "/auditoria/{id_registro}",
     "RF-06: metadatos de auditoría (creador, fechas, versión actual, total versiones)"),
    ("GET",  "/auditoria/historial/{id_registro}",
     "RF-05/09: historial completo con filtros opcionales fecha_desde, fecha_hasta, tipo_cambio"),
    ("GET",  "/auditoria/metadatos/{id_registro}",
     "Metadatos del registro biológico (identificacion_basica + informacion_registro)"),
    ("GET",  "/auditoria/sensibilidad/{id_registro}",
     "RF-04: clasificación PUBLIC / RESTRICTED / CONFIDENTIAL"),
    ("GET",  "/auditoria/registros",
     "Lista paginada de snapshots más recientes de todos los registros"),
    ("GET",  "/auditoria/entradas",
     "Lista paginada de todas las entradas de auditoría"),
    ("GET",  "/aprobacion/{id_registro}",
     "RF-08: estado de aprobación actual del registro"),
    ("POST", "/aprobacion/actualizar",
     "RF-08: actualizar estado de aprobación (genera nueva entrada inmutable)"),
    ("POST", "/dev/simulate",
     "Solo en development: simula un evento Kafka inyectando un registro directamente"),
    ("GET",  "/health",
     "Estado de conexión de MongoDB, caché y Kafka"),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
add_header_row(tbl2, ["Método", "Ruta", "Descripción"], col_widths=[2, 6, 9])
for method, route, desc in endpoints:
    row = tbl2.add_row()
    row.cells[0].text = method
    row.cells[1].text = route
    row.cells[2].text = desc
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].paragraphs[0].runs[0].font.name = "Courier New"
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if method == "GET":
        set_cell_bg(row.cells[0], "DDEEFF")
    else:
        set_cell_bg(row.cells[0], "D5E8D4")

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 7. TESTS
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Tests", 1)

doc.add_heading("7.1 ¿Los tests se ejecutan solos?", 2)

p = doc.add_paragraph(
    "No se ejecutan automáticamente al arrancar el servidor. "
    "Son tests unitarios que se lanzan manualmente desde la línea de comandos "
    "o desde el pipeline de CI/CD. No requieren conexión real a MongoDB, Redis "
    "ni Kafka — todas las dependencias externas están mockeadas con "
)
p.runs[0].font.size = Pt(10)
r = p.add_run("unittest.mock")
r.font.name = "Courier New"
r.font.size = Pt(10)
p.add_run(".").font.size = Pt(10)

section_line(doc)

doc.add_heading("7.2 Cómo ejecutar los tests", 2)

doc.add_heading("Ejecutar todos los tests", 3)
code_paragraph(doc, "pytest")

doc.add_heading("Con salida detallada (ver nombre de cada test)", 3)
code_paragraph(doc, "pytest -v")

doc.add_heading("Con reporte de cobertura en la consola", 3)
code_paragraph(doc, "pytest --cov=services.audit_service --cov-report=term-missing")

doc.add_heading("Generar reporte HTML de cobertura", 3)
code_paragraph(doc, "pytest --cov=services.audit_service --cov-report=html")
code_paragraph(doc, "# El reporte queda en htmlcov/index.html — ábrelo en el navegador")

doc.add_heading("Ejecutar solo un archivo o una clase de tests", 3)
code_paragraph(doc, "pytest tests/test_audit_service.py")
code_paragraph(doc, "pytest tests/test_audit_service.py::TestCreateAuditEntry")
code_paragraph(doc, "pytest tests/test_audit_service.py::TestCreateAuditEntry::test_first_entry_gets_version_1")

note_paragraph(doc,
    "💡  El flag --reload de uvicorn NO ejecuta tests. "
    "pytest y uvicorn son procesos completamente independientes.",
    "D1ECF1")

section_line(doc)

doc.add_heading("7.3 Configuración de pytest", 2)

p = doc.add_paragraph(
    "El archivo pytest.ini en la raíz del proyecto contiene la configuración global:"
)
p.runs[0].font.size = Pt(10)

code_paragraph(doc, "[pytest]")
code_paragraph(doc, "asyncio_mode = auto     # todos los tests async corren sin @pytest.mark.asyncio")
code_paragraph(doc, "testpaths = tests       # pytest busca tests solo en la carpeta tests/")

section_line(doc)

doc.add_heading("7.4 Estructura de los tests existentes", 2)

p = doc.add_paragraph(
    "Actualmente existe un único archivo de tests que cubre el servicio principal "
    "(AuditoriaService), alcanzando el 82% de cobertura requerida por RNF-06."
)
p.runs[0].font.size = Pt(10)

test_classes = [
    ("TestComputeFieldChanges",  "7 tests",
     "Prueba el algoritmo de diferencias entre versiones: sin cambios, campo modificado, campo nuevo, records idénticos."),
    ("TestCreateAuditEntry",    "11 tests",
     "Prueba la creación de entradas: versión inicial (v1), incremento de versión, error sin id_registro, "
     "invalidación de caché (historial y auditoria), almacenamiento de campos, upsert en colección de registros."),
    ("TestGetHistorial",         "3 tests",
     "Prueba la consulta del historial: resultado con entradas, lista vacía, múltiples versiones."),
    ("TestGetLatestSnapshot",    "3 tests",
     "Prueba la obtención del snapshot más reciente: encontrado, no encontrado, sin clave data."),
    ("TestGetAllRecords",        "2 tests",
     "Prueba la paginación de registros: resultado con datos, resultado vacío."),
    ("TestGetAllAuditEntries",   "2 tests",
     "Prueba la paginación de entradas de auditoría: con datos, sin datos."),
    ("TestGetMetadatos",         "4 tests",
     "Prueba los metadatos del registro: ambas secciones presentes, id_registro correcto, "
     "None cuando no existe snapshot, dicts vacíos cuando faltan claves."),
]

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = "Table Grid"
add_header_row(tbl3, ["Clase de test", "Tests", "Qué verifica"], col_widths=[5.5, 2, 9.5])
for cls, count, desc in test_classes:
    row = tbl3.add_row()
    row.cells[0].text = cls
    row.cells[1].text = count
    row.cells[2].text = desc
    row.cells[0].paragraphs[0].runs[0].font.name = "Courier New"
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

section_line(doc)

doc.add_heading("7.5 Tests pendientes de implementar", 2)

pending_tests = [
    ("tests/test_sensitivity_service.py",
     "Función pura sin dependencias externas. Fácil de cubrir al 100%. "
     "Verifica la cadena PUBLIC → RESTRICTED → CONFIDENTIAL desde geolocalizacion y datos_consumo_externo."),
    ("tests/test_kafka_consumer.py",
     "Valida el comportamiento de ECA-02: delay que se duplica en cada fallo, "
     "reset del delay al reconectar, stop() limpio sin excepciones."),
    ("tests/test_approval_service.py",
     "Verifica que get_approval_status retorne el enum correcto y que "
     "update_approval cree una nueva entrada inmutable con tipo_cambio=CAMBIO_ESTADO."),
    ("tests/test_routers_auditoria.py",
     "Tests de integración HTTP con httpx.AsyncClient: "
     "200 en cache miss y hit, 404 para registros inexistentes, filtros de historial."),
]

for filename, desc in pending_tests:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(filename + ": ")
    r.bold = True
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    p.add_run(desc).font.size = Pt(9)

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 8. SIMULAR EVENTOS KAFKA EN DESARROLLO
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Simular eventos Kafka en desarrollo", 1)

p = doc.add_paragraph(
    "Como USE_MOCK_KAFKA=true por defecto, el consumer real no arranca. "
    "Para probar el pipeline completo sin un broker Kafka real, "
    "usa el endpoint de simulación:"
)
p.runs[0].font.size = Pt(10)

code_paragraph(doc, "POST http://localhost:8000/dev/simulate")
code_paragraph(doc, "Content-Type: application/json")
code_paragraph(doc, "")
code_paragraph(doc, "{")
code_paragraph(doc, '  "identificacion_basica": {')
code_paragraph(doc, '    "id_registro": "REG-001",')
code_paragraph(doc, '    "nombre_cientifico": "Panthera onca"')
code_paragraph(doc, "  },")
code_paragraph(doc, '  "informacion_registro": {')
code_paragraph(doc, '    "investigador": "researcher@biohub.org"')
code_paragraph(doc, "  },")
code_paragraph(doc, '  "geolocalizacion": {')
code_parameter = '    "nivel_sensibilidad": "PUBLIC"'
code_paragraph(doc, code_parameter)
code_paragraph(doc, "  }")
code_paragraph(doc, "}")

p = doc.add_paragraph(
    "Este endpoint solo está disponible cuando ENV=development. "
    "Inyecta el registro directamente en audit_service.create_audit_entry, "
    "siguiendo exactamente el mismo flujo que un mensaje real de Kafka."
)
p.runs[0].font.size = Pt(10)

section_line(doc)


# ════════════════════════════════════════════════════════════════════════════════
# 9. DESPLIEGUE
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Despliegue", 1)

p = doc.add_paragraph(
    "El proyecto incluye configuración para Railway (Procfile) y Vercel (vercel.json). "
    "El despliegue en Railway es el recomendado por soportar procesos persistentes "
    "necesarios para el consumer Kafka."
)
p.runs[0].font.size = Pt(10)

deploy = [
    ("Variable de entorno",     "Valor en producción",             "Descripción"),
    ("MONGODB_URI",             "URI de Atlas o instancia propia",  "Obligatorio"),
    ("USE_MOCK_KAFKA",          "false",                            "Activa el consumer real"),
    ("KAFKA_BOOTSTRAP_SERVERS", "host:9092 del broker del Grupo 1", "Dirección del broker"),
    ("REDIS_URL",               "redis://host:6379",               "Opcional, mejora el rendimiento"),
    ("ENV",                     "production",                      "Desactiva /dev/simulate"),
]

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = "Table Grid"
add_header_row(tbl4, ["Variable", "Valor en producción", "Nota"], col_widths=[5, 6, 6])
for i, (var, val, note) in enumerate(deploy[1:]):
    row = tbl4.add_row()
    row.cells[0].text = var
    row.cells[1].text = val
    row.cells[2].text = note
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].font.name = "Courier New"


# ── Save ───────────────────────────────────────────────────────────────────────
output_path = "guia_tecnica_biohub.docx"
doc.save(output_path)
print(f"[OK] Documento generado: {output_path}")
