"""Script de uso único — genera el análisis de brechas ECA vs código real."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def header_row(table, headers, widths=None, bg="1F4E79"):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        set_cell_bg(cell, bg)
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Cm(w)


VERDE   = "E2EFDA"
AMARILLO= "FFEB9C"
ROJO    = "FFC7CE"

def veredicto_bg(v):
    if v.startswith("✅"): return VERDE
    if v.startswith("⚠️"): return AMARILLO
    return ROJO


def data_row(table, values, last_left=True):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        align = WD_ALIGN_PARAGRAPH.LEFT if (last_left and i == len(values)-2) else WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].alignment = align
    return row


def eca_table(doc, title, subtitle, badge_text, badge_bg, rows_data):
    """
    rows_data: list of (criterio, evidencia_rubric, realidad_codigo, veredicto)
    """
    # ECA heading
    h = doc.add_heading(title, 2)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph(subtitle)
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    # badge
    badge_p = doc.add_paragraph()
    badge_run = badge_p.add_run(f"  Estado general: {badge_text}  ")
    badge_run.bold = True
    badge_run.font.size = Pt(10)
    badge_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if badge_bg in ("C00000", "385723") else RGBColor(0x00, 0x00, 0x00)
    # color background via XML on paragraph
    pPr = badge_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), badge_bg)
    pPr.append(shd)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    header_row(tbl,
               ["Criterio del SRS", "Evidencia que pide el rubric", "Realidad en el código", "Veredicto"],
               widths=[3.5, 4.5, 5.5, 2.5])

    for criterio, evidencia, realidad, veredicto in rows_data:
        row = tbl.add_row()
        vals = [criterio, evidencia, realidad, veredicto]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                run.bold = True
            if i == 3:
                set_cell_bg(cell, veredicto_bg(veredicto))

    doc.add_paragraph()


# ── document ───────────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Title
title = doc.add_heading("BioHub — Análisis de Brechas: Rubric vs Código Real", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph("Comparación criterio a criterio entre el rubric de evaluación y la implementación actual")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(11)
sub.runs[0].bold = True

date_p = doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(9)
date_p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ── Resumen ejecutivo ──────────────────────────────────────────────────────────
doc.add_heading("Resumen ejecutivo", 1)

p = doc.add_paragraph(
    "El rubric de evaluación marca ECA-01 y ECA-03 como 'Cumplido' y ECA-02 como 'Parcial'. "
    "Al revisar el código real criterio por criterio, el estado verdadero es diferente:"
)
p.runs[0].font.size = Pt(10)

summary = [
    ("ECA-01 — Rendimiento bajo consulta masiva",
     "Cumplido (rubric)",  "Parcial (real)",
     "Consumer no es proceso independiente. Sin pruebas de carga. Retorna 500, no 503 ante falla Redis."),
    ("ECA-02 — Confiabilidad ante falla Kafka",
     "Parcial (rubric)",   "Más incompleto que parcial",
     "Faltan: session_timeout_ms, enable_auto_commit=False, DLQ y logs estructurados."),
    ("ECA-03 — Integridad ante modificación",
     "Cumplido (rubric)",  "Parcial (real)",
     "Falta el middleware que registra en log de seguridad los intentos de PUT/DELETE."),
]

tbl0 = doc.add_table(rows=1, cols=4)
tbl0.style = "Table Grid"
header_row(tbl0, ["Escenario", "Rubric dice", "Realidad", "Brechas principales"],
           widths=[4.5, 2.5, 3, 7])

for eca, rubric_val, real_val, brechas in summary:
    row = tbl0.add_row()
    for i, val in enumerate([eca, rubric_val, real_val, brechas]):
        cell = row.cells[i]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            run.bold = True
        if i == 1:
            set_cell_bg(cell, VERDE)
        if i == 2:
            if "incompleto" in val:
                set_cell_bg(cell, ROJO)
            else:
                set_cell_bg(cell, AMARILLO)

doc.add_paragraph()

# ── Leyenda ────────────────────────────────────────────────────────────────────
leyenda_p = doc.add_paragraph()
leyenda_p.add_run("Leyenda: ").bold = True
leyenda_p.runs[0].font.size = Pt(9)
leyenda_p.add_run("  ✅ Cumple  ").font.size = Pt(9)
leyenda_p.add_run("  ⚠️ Parcial  ").font.size = Pt(9)
leyenda_p.add_run("  ❌ No cumple  ").font.size = Pt(9)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
# ECA-01
# ════════════════════════════════════════════════════════════════════════════════
eca_table(doc,
    "ECA-01 — Rendimiento bajo consulta masiva",
    "Situación: 500 peticiones simultáneas al mismo registro en 30 segundos.",
    "⚠️  PARCIAL  (rubric lo marca Cumplido)",
    "FFEB9C",
    [
        ("95% peticiones\n< 200 ms",
         "cache-first en RedisRepository.get()\nantes de MongoDB",
         "cache_get() existe antes de ir a MongoDB ✓\nPero NO hay pruebas de carga que\ncertifiquen el objetivo de 200 ms.",
         "⚠️ Parcial"),
        ("Invalidación al\nregistrar cambio",
         "invalidar(id) llamado en\nAuditoriaService.registrar_cambio()",
         "cache_delete('historial:{id}') y\ncache_delete('auditoria:{id}')\nen create_audit_entry ✓",
         "✅ Cumple"),
        ("Escritura Kafka\n< 300 ms en pico",
         "Consumer Worker como proceso\nindependiente del servidor FastAPI",
         "El consumer es asyncio.create_task()\ndentro del MISMO proceso FastAPI.\nNO es un proceso independiente.",
         "❌ No cumple"),
        ("0 errores 5xx\nbajo carga",
         "try/except en cada endpoint —\nretorna 503 controlado si Redis cae",
         "try/except existe en todos los endpoints ✓\nPero retorna HTTP 500, no 503.\nSin manejo específico de falla Redis.",
         "⚠️ Parcial"),
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
# ECA-02
# ════════════════════════════════════════════════════════════════════════════════
eca_table(doc,
    "ECA-02 — Confiabilidad ante falla del broker Kafka",
    "Situación: broker Kafka inaccesible durante 5 minutos.",
    "❌  MÁS INCOMPLETO QUE PARCIAL  (rubric lo marca Parcial)",
    "C00000",
    [
        ("Detección de falla\n< 10 s",
         "session_timeout_ms=10000\nen AIOKafkaConsumer",
         "AIOKafkaConsumer NO tiene\nsession_timeout_ms configurado.\nUsa el valor por defecto del broker.",
         "❌ No cumple"),
        ("Backoff exponencial\n+ logs",
         "asyncio.sleep(2**intento) con\nmax_retries=5 en loop consumer",
         "Backoff existe: delay = min(delay*2, 60) ✓\nPero sin max_retries (reintenta infinito).\nSin logs estructurados, solo print().",
         "⚠️ Parcial"),
        ("0 pérdidas de\nmensajes tras reconexión",
         "enable_auto_commit=False —\ncommit manual tras persistir en MongoDB",
         "AIOKafkaConsumer usa enable_auto_commit\n=True por defecto (nunca se configuró).\nSi cae entre recibir y persistir,\nel mensaje se PIERDE.",
         "❌ No cumple"),
        ("Dead letter queue\npara fallidos",
         "Eventos fallidos en logs estructurados\n— sin colección DLQ aún",
         "Solo print('[!] Error processing\nmessage...') ante fallos.\nSin DLQ, sin logs estructurados.",
         "❌ No cumple"),
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
# ECA-03
# ════════════════════════════════════════════════════════════════════════════════
eca_table(doc,
    "ECA-03 — Integridad ante intento de modificación del historial",
    "Situación: alguien intenta PUT o DELETE sobre el historial.",
    "⚠️  PARCIAL  (rubric lo marca Cumplido)",
    "FFEB9C",
    [
        ("100% PUT/DELETE\nrechazados",
         "Router solo define GET —\nFastAPI retorna 405 automáticamente",
         "Solo existen endpoints GET en\nrouters/auditoria.py ✓\nFastAPI retorna 405 automáticamente.",
         "✅ Cumple"),
        ("Rechazo < 100 ms",
         "Rechazo en routing layer —\nno llega al servicio ni a BD",
         "FastAPI rechaza en la capa de routing\nsin tocar servicios ni MongoDB ✓",
         "✅ Cumple"),
        ("Registro en log\nde seguridad",
         "Middleware loguea método, ruta,\nIP y timestamp de cada request",
         "NO existe ningún middleware de seguridad\nen el proyecto. Los intentos\nrechazados no quedan registrados.",
         "❌ No cumple"),
        ("MongoDB sin\nUPDATE/DELETE",
         "MongoRepository solo expone\ninsertar() y listar*()",
         "La colección audit_entries solo usa\ninsert_one() y find() ✓\n(Sin clase Repository, lógica directa.)",
         "✅ Cumple"),
    ]
)

# ════════════════════════════════════════════════════════════════════════════════
# Plan de acción
# ════════════════════════════════════════════════════════════════════════════════
doc.add_heading("Plan de acción para cerrar las brechas", 1)

acciones = [
    ("1", "enable_auto_commit=False en el consumer",
     "CRÍTICO — evita pérdida real de mensajes Kafka",
     "kafka_service/consumer.py",
     "Agregar enable_auto_commit=False y llamar await consumer.commit() después de que create_audit_entry persista exitosamente en MongoDB."),
    ("2", "session_timeout_ms=10000 en AIOKafkaConsumer",
     "Alta — garantiza detección de falla en < 10 s",
     "kafka_service/consumer.py",
     "Agregar el parámetro session_timeout_ms=10000 (y opcionalmente request_timeout_ms=11000) en la instanciación del AIOKafkaConsumer."),
    ("3", "Middleware de log de seguridad",
     "Alta — cierra ECA-03 completamente",
     "main.py",
     "Agregar @app.middleware('http') que registre método, ruta, IP (request.client.host) y timestamp para cualquier petición con método PUT o DELETE."),
    ("4", "Proceso independiente para el consumer (worker.py)",
     "Media — cierra criterio de separación de ECA-01",
     "worker.py (nuevo)",
     "Crear worker.py con su propio loop asyncio que inicie solo el consumer Kafka sin FastAPI. Actualizar Procfile para Railway con dos procesos: web y worker."),
    ("5", "Logging estructurado (reemplazar print)",
     "Media — resuelve RNF-08 y RNF-05",
     "Todos los módulos",
     "Reemplazar print() por logging.getLogger() con formato JSON que incluya nivel, timestamp, id_registro y resultado. Evita exponer snapshots en trazas de error."),
]

tbl_acc = doc.add_table(rows=1, cols=5)
tbl_acc.style = "Table Grid"
header_row(tbl_acc,
           ["#", "Acción", "Impacto", "Archivo", "Qué hacer"],
           widths=[0.6, 4, 3, 3.5, 6])

impact_colors = {
    "CRÍTICO": "FFC7CE",
    "Alta": "FFEB9C",
    "Media": "DDEBF7",
}

for num, accion, impacto, archivo, que_hacer in acciones:
    row = tbl_acc.add_row()
    vals = [num, accion, impacto, archivo, que_hacer]
    for i, val in enumerate(vals):
        cell = row.cells[i]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
        if i == 0:
            run.bold = True
        if i == 2:
            for key, color in impact_colors.items():
                if key in val:
                    set_cell_bg(cell, color)
                    break

doc.add_paragraph()

# ── Save ───────────────────────────────────────────────────────────────────────
output = "analisis_brechas_eca.docx"
doc.save(output)
print(f"[OK] Documento generado: {output}")
