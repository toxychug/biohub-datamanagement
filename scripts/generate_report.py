"""Script de uso único para generar el reporte Word de estado del proyecto BioHub."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_header_row(table, headers: list, bg: str = "1F4E79"):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(cell, bg)


def add_data_row(table, values: list, bg: str = None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 3 else WD_ALIGN_PARAGRAPH.LEFT
        run = para.runs[0]
        run.font.size = Pt(9)
        if bg:
            set_cell_bg(cell, bg)
    return row


def estado_color(estado: str) -> str:
    if estado.startswith("✅"):
        return "E2EFDA"
    if estado.startswith("⚠️"):
        return "FFEB9C"
    if estado.startswith("❌"):
        return "FFC7CE"
    if estado.startswith("❓"):
        return "DDEBF7"
    return "FFFFFF"


# ── document ───────────────────────────────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("BioHub — Microservicio de Gestión de Cambios", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph("Reporte de Estado de Requerimientos y Escenarios de Calidad")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(12)

date_p = doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(9)
date_p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ── Summary boxes ──────────────────────────────────────────────────────────────
doc.add_heading("Resumen Ejecutivo", 1)

summary_data = [
    ("Requerimientos Funcionales",    "10 / 10", "E2EFDA"),
    ("Requerimientos No Funcionales", "6 / 8",   "FFEB9C"),
    ("Escenarios de Calidad",         "2 activos / 1 parcial", "FFEB9C"),
    ("Cobertura de pruebas",          "82 % en AuditoriaService", "E2EFDA"),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(tbl, ["Área", "Estado"])
for label, value, bg in summary_data:
    row = tbl.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ── RF table ───────────────────────────────────────────────────────────────────
doc.add_heading("Requerimientos Funcionales", 1)

rf_data = [
    ("RF-01", "Consumo de eventos Kafka",         "✅ Cumplido",
     "Consumer con reconexión automática y backoff exponencial (ECA-02)"),
    ("RF-02", "Registro de auditoría",            "✅ Cumplido",
     "create_audit_entry registra campo, valor anterior/nuevo, motivo, usuario, timestamp, IP"),
    ("RF-03", "Gestión de versiones",             "✅ Cumplido",
     "Auto-incremento de versión en cada inserción. El conteo histórico se mantiene en MongoDB"),
    ("RF-04", "Clasificación de sensibilidad",    "✅ Cumplido",
     "sensitivity_service.py clasifica en PUBLIC / RESTRICTED / CONFIDENTIAL"),
    ("RF-05", "Endpoint historial",               "✅ Cumplido",
     "GET /auditoria/historial/{id_registro} — retorna todas las versiones ordenadas"),
    ("RF-06", "Endpoint metadatos de auditoría",  "✅ Cumplido",
     "GET /auditoria/{id_registro} — retorna creador, fechas, versión actual y total de versiones"),
    ("RF-07", "Inmutabilidad del historial",      "✅ Cumplido",
     "La colección de auditoría solo usa insert_one. No existen endpoints PUT/DELETE"),
    ("RF-08", "Flujo de aprobación",              "✅ Cumplido",
     "routers/aprobacion.py gestiona PENDIENTE / EN_REVISION / APROBADO / RECHAZADO"),
    ("RF-09", "Filtrado de historial",            "✅ Cumplido",
     "Query params opcionales: fecha_desde, fecha_hasta, tipo_cambio (CREACION, EDICION, RECLASIFICACION, CAMBIO_ESTADO)"),
    ("RF-10", "Manejo de errores Kafka",          "✅ Cumplido",
     "Backoff exponencial 1 s → 2 s → … → 60 s con reconexión automática al broker"),
]

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.style = "Table Grid"
add_header_row(tbl2, ["ID", "Nombre", "Estado", "Notas"])
tbl2.columns[0].width = Cm(1.5)
tbl2.columns[1].width = Cm(4.5)
tbl2.columns[2].width = Cm(2.5)
tbl2.columns[3].width = Cm(8.5)

for rf_id, name, estado, notes in rf_data:
    bg = estado_color(estado)
    row = tbl2.add_row()
    for i, val in enumerate([rf_id, name, estado, notes]):
        cell = row.cells[i]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        if i == 2:
            set_cell_bg(cell, bg)

doc.add_paragraph()

# ── RNF table ──────────────────────────────────────────────────────────────────
doc.add_heading("Requerimientos No Funcionales", 1)

rnf_data = [
    ("RNF-01", "Disponibilidad 99%",             "✅ Cumplido",
     "Consumer corre como asyncio.create_task independiente; una caída de Kafka no afecta la API REST"),
    ("RNF-02", "Rendimiento <500 ms",            "❓ Sin verificar",
     "La caché ayuda, pero no existen pruebas de carga que certifiquen el objetivo de 200 ms del ECA-01"),
    ("RNF-03", "Caché Redis",                    "✅ Cumplido",
     "Redis con fallback a memoria en-proceso. Invalidación al insertar auditoría y al aprobar"),
    ("RNF-04", "Integridad de datos",            "✅ Cumplido",
     "Colección de auditoría es append-only por diseño de código; no existe lógica de modificación"),
    ("RNF-05", "Seguridad en tránsito",          "⚠️ Parcial",
     "HTTPS depende del deployment (Railway/Vercel). Los print de error pueden exponer snapshots en trazas"),
    ("RNF-06", "Cobertura ≥70% AuditoriaService","✅ Cumplido",
     "82% de cobertura en services/audit_service.py. 31 tests pytest pasando"),
    ("RNF-07", "Escalabilidad del consumer",     "⚠️ Parcial",
     "group_id='biohub-change-management' correcto para múltiples instancias, pero consumer y API comparten el mismo proceso"),
    ("RNF-08", "Observabilidad / Logs",          "❌ Pendiente",
     "Solo existen print básicos. El SRS exige logs estructurados con nivel, timestamp, id_registro y resultado"),
]

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = "Table Grid"
add_header_row(tbl3, ["ID", "Nombre", "Estado", "Notas"])

for rnf_id, name, estado, notes in rnf_data:
    bg = estado_color(estado)
    row = tbl3.add_row()
    for i, val in enumerate([rnf_id, name, estado, notes]):
        cell = row.cells[i]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        if i == 2:
            set_cell_bg(cell, bg)

doc.add_paragraph()

# ── ECA table ──────────────────────────────────────────────────────────────────
doc.add_heading("Escenarios de Calidad", 1)

eca_data = [
    ("ECA-01", "Rendimiento bajo consulta masiva (Caché)",
     "✅ Activo",
     "Caché Redis sirve peticiones repetidas sin ir a MongoDB. Invalidación ocurre en cada nueva entrada de auditoría y en cada cambio de aprobación. Resultados filtrados no se cachean."),
    ("ECA-02", "Confiabilidad ante falla del broker Kafka",
     "✅ Activo",
     "_consume_with_retry reconecta automáticamente con backoff exponencial 1 s → 60 s. La API REST permanece operativa durante la interrupción. Al restaurar el broker, el consumer retoma desde el último offset."),
    ("ECA-03", "Integridad ante intento de modificación del historial",
     "⚠️ Parcial",
     "FastAPI retorna 405 automáticamente para PUT/DELETE no definidos — el rechazo funciona al 100%. Pendiente: middleware que registre en log de seguridad el usuario, IP y timestamp de cada intento."),
]

tbl4 = doc.add_table(rows=1, cols=4)
tbl4.style = "Table Grid"
add_header_row(tbl4, ["ID", "Escenario", "Estado", "Detalle"])
tbl4.columns[0].width = Cm(1.5)
tbl4.columns[1].width = Cm(5.0)
tbl4.columns[2].width = Cm(2.5)
tbl4.columns[3].width = Cm(8.0)

for eca_id, name, estado, detail in eca_data:
    bg = estado_color(estado)
    row = tbl4.add_row()
    for i, val in enumerate([eca_id, name, estado, detail]):
        cell = row.cells[i]
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        if i == 2:
            set_cell_bg(cell, bg)

doc.add_paragraph()

# ── Pending section ────────────────────────────────────────────────────────────
doc.add_heading("Trabajo pendiente", 1)

pending = [
    ("RNF-08 + RNF-05",
     "Logging estructurado",
     "Reemplazar los print por Python logging con formato JSON (nivel, timestamp, id_registro, resultado). "
     "Al controlar qué campos se loguean se resuelve simultáneamente la restricción de no exponer datos sensibles en logs (RNF-05)."),
    ("RNF-07",
     "Separación consumer / API en procesos independientes",
     "Crear un worker.py independiente que ejecute únicamente el consumer Kafka, separado del proceso FastAPI. "
     "Esto permite escalar horizontalmente cada componente de forma independiente mediante docker-compose o Railway."),
    ("ECA-03 (incompleto)",
     "Log de seguridad para intentos de modificación",
     "Agregar un middleware o exception handler que registre en log los intentos de PUT/DELETE sobre "
     "/auditoria/historial/, capturando usuario, IP y timestamp de cada intento rechazado."),
]

for label, title_text, description in pending:
    p = doc.add_paragraph(style="List Bullet")
    run_label = p.add_run(f"[{label}] {title_text}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_desc = p.add_run(description)
    run_desc.font.size = Pt(10)

doc.add_paragraph()

# ── Legend ─────────────────────────────────────────────────────────────────────
doc.add_heading("Leyenda", 2)

legend_items = [
    ("E2EFDA", "✅ Cumplido     — Requerimiento implementado y verificado"),
    ("FFEB9C", "⚠️ Parcial      — Implementación incompleta o con brechas menores"),
    ("FFC7CE", "❌ Pendiente    — No implementado"),
    ("DDEBF7", "❓ Sin verificar — Implementado pero sin pruebas que lo certifiquen"),
]

tbl5 = doc.add_table(rows=len(legend_items), cols=1)
tbl5.style = "Table Grid"
for i, (bg, text) in enumerate(legend_items):
    cell = tbl5.rows[i].cells[0]
    cell.text = text
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(cell, bg)

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = "reporte_estado_biohub.docx"
doc.save(output_path)
print(f"[OK] Documento generado: {output_path}")
